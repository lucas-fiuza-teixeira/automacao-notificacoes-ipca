from docx import Document
from docx.shared import Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

from utils.docx_utils import aplicar_bordas
from utils.formatacao import limpar_nome_arquivo
from services.calculo_service import calcular

def gerar_notificacao(grupo, dados_cliente, template_path, output_path):
    razao, cnpj, contrato, email = dados_cliente

    doc = Document(template_path)
    total_geral = 0
    memoria = []

    p_tabela = None
    for p in doc.paragraphs:
        if "[TABELA INSERIDA AUTOMATICAMENTE]" in p.text:
            p_tabela = p
            break

    tabela = doc.add_table(rows=1, cols=7)

    if p_tabela:
        p_tabela._element.addnext(tabela._element)
        p_tabela.text = p_tabela.text.replace("[TABELA INSERIDA AUTOMATICAMENTE]", "")

    aplicar_bordas(tabela)

    titulos = ["NF", "Vencimento", "Base", "Corrigido", "Multa", "Juros", "Total"]
    for i, t in enumerate(titulos):
        tabela.rows[0].cells[i].text = t

    for _, r in grupo.iterrows():
        venc = r['data_vencimento']

        base, corr, multa, juros, total = calcular(r['valor'], venc)
        total_geral += total

        memoria.append({
            "CNPJ": cnpj,
            "NF": r['numero_nf'],
            "Total": total
        })

        linha = tabela.add_row().cells
        linha[0].text = str(r['numero_nf'])
        linha[1].text = venc.strftime("%d/%m/%Y")
        linha[2].text = f"{base:.2f}"
        linha[3].text = f"{corr:.2f}"
        linha[4].text = f"{multa:.2f}"
        linha[5].text = f"{juros:.2f}"
        linha[6].text = f"{total:.2f}"

    nome = limpar_nome_arquivo(razao)
    caminho = f"{output_path}/notificacao_{nome}.docx"
    doc.save(caminho)

    return total_geral, memoria